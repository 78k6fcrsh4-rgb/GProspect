"""
agent/digest.py
---------------
Weekly digest assembly + rendering (Phase 1b).

Three layers:

  1. build_digest_payload(org, opportunities, current_profile) -> dict
     Pure function. Picks the top-5 opportunities, builds the 30-day
     deadline calendar, and returns a JSON-able payload that's easy to
     test independently of the renderers.

  2. render_docx(payload) -> bytes
     Renders the payload as a Microsoft Word document via python-docx.

  3. render_pdf(payload) -> bytes
     Renders the same payload as a PDF via reportlab.

The two renderers consume the same payload so visual differences come
from the renderer choice, not from divergent content. PRD called for
"top 5 opportunities + 30-day deadline calendar + landscape changes" —
landscape changes (week-over-week new funders, dropped opportunities)
is deferred until Phase 2 when we have a discovery cycle producing
weekly snapshots.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Any

from agent.opportunities import classify_deadline, parse_days_remaining

log = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Payload assembly
# ─────────────────────────────────────────────────────────────────────────────

def build_digest_payload(
    org_display_name: str,
    opportunities:    list[dict],
    profile_payload:  dict | None = None,
    narratives:       dict[str, str] | None = None,
    generated_at:     datetime | None = None,
) -> dict:
    """
    Build the JSON payload the renderers consume.

    Args:
        org_display_name: Human-readable org name for the header.
        opportunities:    Enriched opportunity rows (from /opportunities).
                          Must include funder_name, program_name, score_final,
                          application_deadline, days_remaining, award_range,
                          next_action, application_url, opp_key.
        profile_payload:  The current OrgProfile payload, for one-line org
                          context in the header. Optional.
        narratives:       Mapping of opp_key → conversational narrative
                          markdown. Used to inline the "why this fits"
                          paragraph in the top-5 section. Anything missing
                          falls back to the next_action field.
        generated_at:     Override generation timestamp (defaults to now).

    Returns:
        Dict with: org, generated_at, top_five, deadline_calendar,
        and metadata for the renderers.
    """
    narratives = narratives or {}
    when = generated_at or datetime.now(timezone.utc)

    # Top 5 by score, dropping anything without a score
    scored = [
        o for o in opportunities
        if o.get("score_final") is not None
    ]
    scored.sort(key=lambda o: o.get("score_final") or 0, reverse=True)
    top_five = scored[:5]

    # Deadline calendar — next 30 days, sorted by days_remaining ascending
    deadline_calendar: list[dict] = []
    for o in opportunities:
        days = parse_days_remaining(o.get("days_remaining"))
        if days is None or days < 0 or days > 30:
            continue
        deadline_calendar.append({
            "funder_name":          o.get("funder_name"),
            "program_name":         o.get("program_name"),
            "application_deadline": o.get("application_deadline"),
            "days_remaining":       days,
            "bucket":               classify_deadline(days),
            "application_url":      o.get("application_url"),
        })
    deadline_calendar.sort(key=lambda d: d["days_remaining"])

    # Inline a short org-context line if we have a profile.
    org_context = _build_org_context(profile_payload)

    return {
        "org": {
            "display_name": org_display_name,
            "context":      org_context,
        },
        "generated_at": when.isoformat(),
        "top_five": [
            _enriched_top_row(o, narratives.get(o.get("opp_key", ""), ""))
            for o in top_five
        ],
        "deadline_calendar": deadline_calendar,
        "total_opportunities": len(opportunities),
    }


def _build_org_context(profile_payload: dict | None) -> str:
    """A short single-line summary of the org for the digest header."""
    if not profile_payload:
        return ""
    geo = profile_payload.get("geography") or {}
    city = geo.get("city") or ""
    state = geo.get("state") or ""
    programs = profile_payload.get("program_areas") or []
    parts = []
    if city and state:
        parts.append(f"{city}, {state}")
    elif state:
        parts.append(state)
    if programs:
        parts.append(", ".join(programs[:3]))
    return " · ".join(parts)


def _enriched_top_row(opp: dict, narrative_md: str) -> dict:
    """Trim an opportunity row down to digest-friendly fields."""
    days = parse_days_remaining(opp.get("days_remaining"))
    return {
        "rank":                 opp.get("rank"),
        "funder_name":          opp.get("funder_name"),
        "program_name":         opp.get("program_name"),
        "score_final":          opp.get("score_final"),
        "application_deadline": opp.get("application_deadline"),
        "days_remaining":       days,
        "bucket":               classify_deadline(days),
        "award_range":          opp.get("award_range"),
        "next_action":          opp.get("next_action"),
        "application_url":      opp.get("application_url"),
        "narrative_md":         narrative_md,
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX rendering
# ─────────────────────────────────────────────────────────────────────────────

def render_docx(payload: dict) -> bytes:
    """Render the digest payload as a .docx and return the bytes."""
    from docx              import Document
    from docx.shared       import Pt, RGBColor

    doc = Document()

    # ── Header ────────────────────────────────────────────────────────────────
    title = doc.add_heading("GProspect — Weekly Digest", level=0)
    doc.add_paragraph(payload["org"]["display_name"])
    if payload["org"].get("context"):
        ctx_p = doc.add_paragraph(payload["org"]["context"])
        for run in ctx_p.runs:
            run.italic = True
    doc.add_paragraph(f"Generated {payload['generated_at']}")

    # ── Top 5 ─────────────────────────────────────────────────────────────────
    doc.add_heading("Top 5 opportunities", level=1)
    if not payload["top_five"]:
        doc.add_paragraph("No scored opportunities yet. Trigger a run from the portal.")
    for o in payload["top_five"]:
        h = doc.add_heading(
            f"{o.get('rank') or ''}. {o['funder_name']} — {o['program_name']}",
            level = 2,
        )
        meta_bits = []
        if o.get("score_final") is not None:
            meta_bits.append(f"Score: {o['score_final']:.2f}")
        if o.get("application_deadline"):
            bucket = (o.get("bucket") or "").upper() or ""
            badge  = f" [{bucket}]" if bucket else ""
            meta_bits.append(f"Deadline: {o['application_deadline']}{badge}")
        if o.get("award_range"):
            meta_bits.append(f"Award: {o['award_range']}")
        if meta_bits:
            doc.add_paragraph(" · ".join(meta_bits))

        if o.get("narrative_md"):
            doc.add_paragraph(o["narrative_md"])
        elif o.get("next_action"):
            doc.add_paragraph(f"Next action: {o['next_action']}")

        if o.get("application_url"):
            doc.add_paragraph(f"Apply: {o['application_url']}")

    # ── Deadline calendar ─────────────────────────────────────────────────────
    doc.add_heading("Deadlines in the next 30 days", level=1)
    if not payload["deadline_calendar"]:
        doc.add_paragraph("No deadlines in the next 30 days.")
    else:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Days"
        hdr[1].text = "Deadline"
        hdr[2].text = "Opportunity"
        for d in payload["deadline_calendar"]:
            row = table.add_row().cells
            row[0].text = str(d["days_remaining"])
            row[1].text = d.get("application_deadline") or ""
            row[2].text = f"{d.get('funder_name') or ''} — {d.get('program_name') or ''}"

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"Generated by GProspect over {payload['total_opportunities']} "
        f"opportunities in your pipeline."
    )
    for run in footer.runs:
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PDF rendering
# ─────────────────────────────────────────────────────────────────────────────

def render_pdf(payload: dict) -> bytes:
    """Render the digest payload as a PDF and return the bytes."""
    from reportlab.lib.pagesizes  import letter
    from reportlab.lib.styles     import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units      import inch
    from reportlab.lib            import colors
    from reportlab.platypus       import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )

    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=letter,
                               leftMargin=0.75*inch, rightMargin=0.75*inch,
                               topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    body   = styles["BodyText"]
    h1     = styles["Heading1"]
    h2     = styles["Heading2"]
    h3     = styles["Heading3"]
    small  = ParagraphStyle("small", parent=body, fontSize=9, textColor=colors.grey)

    story = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph("GProspect — Weekly Digest", h1))
    story.append(Paragraph(payload["org"]["display_name"], h2))
    if payload["org"].get("context"):
        story.append(Paragraph(f"<i>{payload['org']['context']}</i>", body))
    story.append(Paragraph(f"Generated {payload['generated_at']}", small))
    story.append(Spacer(1, 0.2*inch))

    # ── Top 5 ─────────────────────────────────────────────────────────────────
    story.append(Paragraph("Top 5 opportunities", h2))
    if not payload["top_five"]:
        story.append(Paragraph("No scored opportunities yet.", body))
    for o in payload["top_five"]:
        title = f"{o.get('rank') or ''}. {o['funder_name']} — {o['program_name']}"
        story.append(Paragraph(title, h3))

        meta_bits = []
        if o.get("score_final") is not None:
            meta_bits.append(f"Score: {o['score_final']:.2f}")
        if o.get("application_deadline"):
            bucket = (o.get("bucket") or "").upper() or ""
            badge  = f" [{bucket}]" if bucket else ""
            meta_bits.append(f"Deadline: {o['application_deadline']}{badge}")
        if o.get("award_range"):
            meta_bits.append(f"Award: {o['award_range']}")
        if meta_bits:
            story.append(Paragraph(" · ".join(meta_bits), small))

        if o.get("narrative_md"):
            # ReportLab's Paragraph accepts a tiny HTML subset — escape angle brackets.
            text = (o["narrative_md"].replace("<", "&lt;").replace(">", "&gt;"))
            story.append(Paragraph(text, body))
        elif o.get("next_action"):
            story.append(Paragraph(f"Next action: {o['next_action']}", body))

        if o.get("application_url"):
            url = o["application_url"]
            story.append(Paragraph(
                f'Apply: <a href="{url}" color="blue">{url}</a>', body,
            ))
        story.append(Spacer(1, 0.15*inch))

    # ── Deadline calendar ─────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Deadlines in the next 30 days", h2))
    if not payload["deadline_calendar"]:
        story.append(Paragraph("No deadlines in the next 30 days.", body))
    else:
        data = [["Days", "Deadline", "Opportunity"]]
        for d in payload["deadline_calendar"]:
            data.append([
                str(d["days_remaining"]),
                d.get("application_deadline") or "",
                f"{d.get('funder_name') or ''} — {d.get('program_name') or ''}",
            ])
        table = Table(data, colWidths=[0.6*inch, 1.4*inch, 4.5*inch])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(
        f"Generated by GProspect over {payload['total_opportunities']} "
        f"opportunities in your pipeline.",
        small,
    ))

    doc.build(story)
    return buf.getvalue()
